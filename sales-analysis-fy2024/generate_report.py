from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

doc = Document()

# ── PAGE MARGINS ───────────────────────────────────────────────
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1.2)
section.right_margin = Inches(1.2)

# ── HELPER FUNCTIONS ───────────────────────────────────────────
def add_heading(doc, text, level=1, color=RGBColor(0x1F, 0x49, 0x7D)):
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in heading.runs:
        run.font.color.rgb = color
    return heading

def add_paragraph(doc, text, bold=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return p

def add_bullet(doc, text, bold_part=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_part and text.startswith(bold_part):
        run1 = p.add_run(bold_part)
        run1.bold = True
        run1.font.size = Pt(11)
        run2 = p.add_run(text[len(bold_part):])
        run2.font.size = Pt(11)
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
    return p


# ── TITLE PAGE ─────────────────────────────────────────────────
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('FY2024 Sales Performance Analysis Report')
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = subtitle.add_run('B2B Office Supplies | Analytics Division')
run2.font.size = Pt(13)
run2.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = meta.add_run('Prepared by: Rishab (Data Analyst Intern)\nPeriod: FY2024 (January – December 2024)\nDataset: 1,000 transactions | 8 Cities | 3 Categories')
run3.font.size = Pt(11)

doc.add_page_break()


# ── SECTION 1: EXECUTIVE SUMMARY ───────────────────────────────
add_heading(doc, '1. Executive Summary')
add_paragraph(doc,
    'This report presents a comprehensive analysis of FY2024 sales data for a B2B office '
    'supplies company operating across 8 major Indian cities. The analysis covers revenue '
    'trends, product performance, regional insights, channel effectiveness, and customer '
    'satisfaction based on 1,000 transaction records.')
doc.add_paragraph()

# KPI Summary Table
add_paragraph(doc, 'Key Performance Indicators', bold=True)
table = doc.add_table(rows=4, cols=2)
table.style = 'Table Grid'
kpis = [
    ('Total Revenue FY2024', '₹5.08 Crore'),
    ('Total Orders', '1,000'),
    ('Average Order Value', '₹50,803'),
    ('Overall Customer Rating', '3.01 / 5'),
]
for i, (metric, value) in enumerate(kpis):
    row = table.rows[i]
    row.cells[0].text = metric
    row.cells[1].text = value
    for cell in row.cells:
        cell.paragraphs[0].runs[0].font.size = Pt(11)
doc.add_paragraph()


# ── SECTION 2: REVENUE & SALES TRENDS ─────────────────────────
add_heading(doc, '2. Revenue & Sales Trends')

doc.add_picture('monthly_revenue_trend.png', width=Inches(5.5))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
add_paragraph(doc, 'Figure 1: Monthly Revenue Trend FY2024', size=9).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

add_bullet(doc, 'Q4 is the strongest quarter at ₹1.65 crore, driven by an October spike of ₹77.5 lakh — likely festive season (Dussehra/Diwali) demand.', 'Q4 is the strongest quarter')
add_bullet(doc, 'Q3 is the weakest quarter at ₹1.05 crore — July to September shows a consistent dip that needs a targeted sales push next year.', 'Q3 is the weakest quarter')
add_bullet(doc, 'January is the weakest individual month at ₹27 lakh, suggesting slow post-holiday business activity.', 'January is the weakest individual month')
doc.add_paragraph()

doc.add_picture('quarterly_revenue.png', width=Inches(4.5))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
add_paragraph(doc, 'Figure 2: Revenue by Quarter', size=9).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()


# ── SECTION 3: PRODUCT PERFORMANCE ────────────────────────────
add_heading(doc, '3. Product Performance')

add_paragraph(doc, 'Top 5 Products by Revenue', bold=True)
prod_table = doc.add_table(rows=6, cols=2)
prod_table.style = 'Table Grid'
prod_data = [
    ('Product', 'Revenue'),
    ('Laptop Pro 15', '₹23,026,250'),
    ('Standing Desk', '₹8,098,200'),
    ('Monitor 24 inch', '₹6,703,200'),
    ('Office Chair', '₹4,214,400'),
    ('Filing Cabinet', '₹2,359,175'),
]
for i, (prod, rev) in enumerate(prod_data):
    row = prod_table.rows[i]
    row.cells[0].text = prod
    row.cells[1].text = rev
    for cell in row.cells:
        run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(cell.text)
        run.font.size = Pt(11)
        if i == 0:
            run.bold = True
doc.add_paragraph()

add_bullet(doc, 'Laptop Pro 15 alone contributes ₹2.3 crore — 45% of total revenue from one product. This is a significant concentration risk.', 'Laptop Pro 15 alone')
add_bullet(doc, 'Bottom 5 products are all Stationery items with combined revenue under ₹6.5 lakh for the full year.', 'Bottom 5 products')
add_bullet(doc, 'Discount strategy is ineffective: orders with 0% discount generated comparable revenue to orders with 20% discount, indicating discounts are not driving proportional volume uplift.', 'Discount strategy is ineffective')
add_bullet(doc, 'Stationery has the highest return rate at 7.03%, followed by Furniture at 6.77% and Electronics at 4.39%.', 'Stationery has the highest return rate')
doc.add_paragraph()

doc.add_picture('monthly_revenue_trend.png', width=Inches(5.0))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
add_paragraph(doc, 'Figure 3: Top 5 Products by Revenue', size=9).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()


# ── SECTION 4: REGIONAL & CITY ANALYSIS ───────────────────────
add_heading(doc, '4. Regional & City Analysis')

add_bullet(doc, 'South region leads with ₹1.97 crore (39% of total revenue), nearly double the West region at ₹1.22 crore.', 'South region leads')
add_bullet(doc, 'Chennai is the top-performing city at ₹76 lakh. Delhi is the weakest major metro at ₹55 lakh despite its market size.', 'Chennai is the top-performing city')
add_bullet(doc, 'Guwahati (Northeast) performs on par with Kolkata (East) at ₹66.5 lakh each — strong Tier-2 city performance.', 'Guwahati (Northeast) performs')
add_bullet(doc, 'Northeast and East regions are nearly equal in revenue — both underperform relative to South and West.', 'Northeast and East regions')
doc.add_paragraph()

doc.add_picture('city_category_heatmap.png', width=Inches(5.5))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
add_paragraph(doc, 'Figure 4: Revenue Heatmap — City vs Category', size=9).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()


# ── SECTION 5: CHANNEL & CUSTOMER INSIGHTS ────────────────────
add_heading(doc, '5. Channel & Customer Insights')

add_bullet(doc, 'All three sales channels are nearly equal in revenue — Retail Store (₹1.74 cr), Corporate B2B (₹1.72 cr), Online (₹1.62 cr). Healthy diversification with no over-dependence on one channel.', 'All three sales channels')
add_bullet(doc, 'Individual customers have the highest average order value at ₹53,541 — higher than Enterprise (₹51,587) and SME (₹47,268). Bulk discounts in Enterprise deals likely suppress per-order value.', 'Individual customers have the highest')
add_bullet(doc, 'Cash is the most used payment method (213 orders), followed closely by EMI (208). UPI adoption is relatively low at 193 orders for a B2B business.', 'Cash is the most used payment method')
doc.add_paragraph()


# ── SECTION 6: CUSTOMER SATISFACTION ──────────────────────────
add_heading(doc, '6. Customer Satisfaction')

add_bullet(doc, 'Overall average rating is 3.01/5 — mediocre. Significant room for improvement across all categories and cities.', 'Overall average rating is 3.01/5')
add_bullet(doc, 'Bangalore has the lowest city rating at 2.85 — delivery reliability, product quality, or after-sales service issues should be investigated.', 'Bangalore has the lowest city rating')
add_bullet(doc, 'Pune has the highest rating at 3.18, followed by Mumbai at 3.13.', 'Pune has the highest rating')
add_bullet(doc, 'Correlation between discount percentage and customer rating is near zero (0.02) — higher discounts do not lead to better customer satisfaction.', 'Correlation between discount percentage')
doc.add_paragraph()


# ── SECTION 7: RECOMMENDATIONS ────────────────────────────────
add_heading(doc, '7. Recommendations')

add_paragraph(doc, 'Based on the analysis, the following actions are recommended for Q1 FY2025:', bold=False)
doc.add_paragraph()

recs = [
    ('Reduce blanket discounting.',
     ' Discounts are not driving proportional revenue or satisfaction gains. Cap discounts at 10% and run targeted promotions only during Q3 (Jul–Sep) to address the seasonal dip.'),
    ('Reduce dependency on Laptop Pro 15.',
     ' A single product driving 45% of revenue is a business risk. Actively promote Standing Desk and Monitor 24 inch — both are in the top 5 and have growth potential.'),
    ('Invest in the South region and investigate Delhi.',
     ' South already leads — double down with dedicated sales resources. Delhi underperforms for its market size — assign a focused sales rep and run a targeted B2B outreach campaign.'),
    ('Improve customer satisfaction in Bangalore.',
     ' With the lowest rating of 2.85, Bangalore needs a service quality audit — check delivery timelines, return processing speed, and product quality for that city specifically.'),
    ('Re-evaluate the Stationery category.',
     ' Stationery contributes less than 2% of total revenue (₹5.6 lakh) and has the highest return rate. Consider reducing SKUs to only high-margin items and dropping slow movers like Whiteboard Marker Set.'),
]

for i, (bold_part, rest) in enumerate(recs):
    p = doc.add_paragraph(style='List Number')
    run1 = p.add_run(bold_part)
    run1.bold = True
    run1.font.size = Pt(11)
    run2 = p.add_run(rest)
    run2.font.size = Pt(11)

doc.add_paragraph()


# ── SECTION 8: CONCLUSION ──────────────────────────────────────
add_heading(doc, '8. Conclusion')
add_paragraph(doc,
    'FY2024 showed solid overall performance with ₹5.08 crore in revenue across 1,000 orders. '
    'However, the business faces key risks: over-reliance on a single product, an ineffective '
    'discount strategy, geographic revenue concentration, and below-average customer satisfaction. '
    'Addressing these through focused regional expansion, smarter pricing, and service quality '
    'improvements can meaningfully improve FY2025 performance.')
doc.add_paragraph()
add_paragraph(doc, 'Tools Used: Python (Pandas, Matplotlib, Seaborn) | Dataset: 1,000 rows, 18 columns', size=10)


# ── SAVE ───────────────────────────────────────────────────────
doc.save('FY2024_Sales_Analysis_Report.docx')
print("Report saved as FY2024_Sales_Analysis_Report.docx")